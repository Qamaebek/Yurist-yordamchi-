from docs_uzbek import make_doc
import asyncio
import logging
import os
import json
import aiosqlite
from datetime import datetime
from aiogram import Bot, Dispatcher, F
from aiogram.filters import CommandStart, Command
from aiogram.types import Message, CallbackQuery, FSInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton
import httpx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO)

BOT_TOKEN = os.getenv("BOT_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
ADMIN_IDS = list(map(int, os.getenv("ADMIN_IDS", "0").split(",")))
DB = "bot.db"
DOCS_DIR = "docs"
os.makedirs(DOCS_DIR, exist_ok=True)

# ─── TILLAR ───────────────────────────────────────────────────────────────────

TEXTS = {
    "uz": {
        "welcome": "⚖️ <b>Yuridik Yordamchi Botiga Xush Kelibsiz!</b>\n\n📄 Hujjatlar konstruktori\n🤖 AI huquqiy maslahat\n\nXizmatni tanlang 👇",
        "choose_doc": "📂 <b>Hujjat turini tanlang:</b>",
        "info": "ℹ️ <b>Xizmatlar narxi:</b>\n\n• Barcha hujjatlar — 10,000 so'm\n• AI Maslahat — 10,000 so'm",
        "ai_start": "🤖 <b>AI Huquqiy Maslahat</b>\n\nHuquqiy muammoingizni yozing:",
        "thinking": "🤔 Tahlil qilinmoqda...",
        "ai_error": "❌ AI xizmati hozircha mavjud emas.",
        "paid_ok": "✅ <b>To'lov tasdiqlandi! Hujjatingiz tayyor.</b>\n\nWord (.docx) formatida yuborildi.",
        "preparing": "⏳ Hujjat tayyorlanmoqda...",
        "cancel": "❌ Bekor qilish",
        "back": "⬅️ Orqaga",
        "menu_docs": "📄 Hujjat olish",
        "menu_ai": "🤖 AI Maslahat",
        "menu_info": "ℹ️ Ma'lumot",
        "menu_lang": "🌐 Til",
        "pay_btn": "💳 To'lash (Demo)",
        "confirm_btn": "✅ Tasdiqlash (Demo)",
        "admin_title": "👨‍💼 <b>Admin Panel</b>",
        "stats_title": "📊 <b>Statistika</b>",
        "broadcast_ask": "📢 Barcha foydalanuvchilarga yuboriladigan xabarni yozing:",
        "broadcast_done": "✅ {n} ta foydalanuvchiga yuborildi.",
        "no_access": "❌ Ruxsat yo'q.",
        "choose_lang": "🌐 Tilni tanlang:",
        "lang_saved": "✅ Til saqlandi!",
    },
    "ru": {
        "welcome": "⚖️ <b>Добро пожаловать в Юридический Помощник!</b>\n\n📄 Конструктор документов\n🤖 AI юридическая консультация\n\nВыберите услугу 👇",
        "choose_doc": "📂 <b>Выберите тип документа:</b>",
        "info": "ℹ️ <b>Стоимость услуг:</b>\n\n• Все документы — 10,000 сум\n• AI консультация — 10,000 сум",
        "ai_start": "🤖 <b>AI Юридическая Консультация</b>\n\nОпишите вашу юридическую проблему:",
        "thinking": "🤔 Анализируется...",
        "ai_error": "❌ AI сервис временно недоступен.",
        "paid_ok": "✅ <b>Оплата подтверждена! Ваш документ готов.</b>\n\nОтправлен в формате Word (.docx).",
        "preparing": "⏳ Документ готовится...",
        "cancel": "❌ Отмена",
        "back": "⬅️ Назад",
        "menu_docs": "📄 Получить документ",
        "menu_ai": "🤖 AI Консультация",
        "menu_info": "ℹ️ Информация",
        "menu_lang": "🌐 Язык",
        "pay_btn": "💳 Оплатить (Демо)",
        "confirm_btn": "✅ Подтвердить (Демо)",
        "admin_title": "👨‍💼 <b>Панель администратора</b>",
        "stats_title": "📊 <b>Статистика</b>",
        "broadcast_ask": "📢 Напишите сообщение для всех пользователей:",
        "broadcast_done": "✅ Отправлено {n} пользователям.",
        "no_access": "❌ Нет доступа.",
        "choose_lang": "🌐 Выберите язык:",
        "lang_saved": "✅ Язык сохранён!",
    },
    "en": {
        "welcome": "⚖️ <b>Welcome to Legal Assistant Bot!</b>\n\n📄 Document constructor\n🤖 AI legal consultation\n\nChoose a service 👇",
        "choose_doc": "📂 <b>Select document type:</b>",
        "info": "ℹ️ <b>Service prices:</b>\n\n• All documents — 10,000 UZS\n• AI consultation — 10,000 UZS",
        "ai_start": "🤖 <b>AI Legal Consultation</b>\n\nDescribe your legal issue:",
        "thinking": "🤔 Analyzing...",
        "ai_error": "❌ AI service is temporarily unavailable.",
        "paid_ok": "✅ <b>Payment confirmed! Your document is ready.</b>\n\nSent in Word (.docx) format.",
        "preparing": "⏳ Preparing document...",
        "cancel": "❌ Cancel",
        "back": "⬅️ Back",
        "menu_docs": "📄 Get document",
        "menu_ai": "🤖 AI Consultation",
        "menu_info": "ℹ️ Information",
        "menu_lang": "🌐 Language",
        "pay_btn": "💳 Pay (Demo)",
        "confirm_btn": "✅ Confirm (Demo)",
        "admin_title": "👨‍💼 <b>Admin Panel</b>",
        "stats_title": "📊 <b>Statistics</b>",
        "broadcast_ask": "📢 Write a message for all users:",
        "broadcast_done": "✅ Sent to {n} users.",
        "no_access": "❌ Access denied.",
        "choose_lang": "🌐 Choose language:",
        "lang_saved": "✅ Language saved!",
    }
}

DOC_NAMES = {
    "uz": {
        "ijara": "🏠 Ijara shartnomasi",
        "oldi_sotdi": "🤝 Oldi-sotdi shartnomasi",
        "tilxat": "✍️ Tilxat",
        "ishonchnoma": "📋 Ishonchnoma",
        "nikoh": "💍 Nikoh shartnomasi",
        "qarz": "💰 Qarz shartnomasi",
    },
    "ru": {
        "ijara": "🏠 Договор аренды",
        "oldi_sotdi": "🤝 Договор купли-продажи",
        "tilxat": "✍️ Расписка",
        "ishonchnoma": "📋 Доверенность",
        "nikoh": "💍 Брачный договор",
        "qarz": "💰 Договор займа",
    },
    "en": {
        "ijara": "🏠 Lease agreement",
        "oldi_sotdi": "🤝 Sale agreement",
        "tilxat": "✍️ Receipt",
        "ishonchnoma": "📋 Power of attorney",
        "nikoh": "💍 Marriage contract",
        "qarz": "💰 Loan agreement",
    }
}

PRICE = 10000

def t(lang, key):
    return TEXTS.get(lang, TEXTS["uz"]).get(key, key)

def dn(lang, doc):
    return DOC_NAMES.get(lang, DOC_NAMES["uz"]).get(doc, doc)

# ─── KLAVIATURALAR ────────────────────────────────────────────────────────────

def main_menu(lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "menu_docs"), callback_data="menu_docs")],
        [InlineKeyboardButton(text=t(lang, "menu_ai"), callback_data="menu_ai")],
        [InlineKeyboardButton(text=t(lang, "menu_info"), callback_data="menu_info")],
        [InlineKeyboardButton(text=t(lang, "menu_lang"), callback_data="menu_lang")],
    ])

def docs_menu(lang="uz"):
    rows = []
    for key in ["ijara", "oldi_sotdi", "tilxat", "ishonchnoma", "nikoh", "qarz"]:
        rows.append([InlineKeyboardButton(
            text=f"{dn(lang, key)} — 10,000",
            callback_data=f"doc_{key}"
        )])
    rows.append([InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")])
    return InlineKeyboardMarkup(inline_keyboard=rows)

def pay_menu(order_id, lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "pay_btn"), callback_data=f"pay_{order_id}")],
        [InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")],
    ])

def confirm_menu(order_id, lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "confirm_btn"), callback_data=f"confirm_{order_id}")],
        [InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")],
    ])

def cancel_menu(lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text=t(lang, "cancel"), callback_data="back_main")],
    ])

def lang_menu():
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🇺🇿 O'zbek", callback_data="setlang_uz")],
        [InlineKeyboardButton(text="🇷🇺 Русский", callback_data="setlang_ru")],
        [InlineKeyboardButton(text="🇬🇧 English", callback_data="setlang_en")],
    ])

def admin_menu(lang="uz"):
    return InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📊 Statistika", callback_data="admin_stats")],
        [InlineKeyboardButton(text="👥 Foydalanuvchilar", callback_data="admin_users")],
        [InlineKeyboardButton(text="📋 So'nggi buyurtmalar", callback_data="admin_orders")],
        [InlineKeyboardButton(text="📢 Xabar tarqatish", callback_data="admin_broadcast")],
        [InlineKeyboardButton(text=t(lang, "back"), callback_data="back_main")],
    ])

# ─── MA'LUMOTLAR BAZASI ───────────────────────────────────────────────────────

async def init_db():
    async with aiosqlite.connect(DB) as db:
        await db.execute("""CREATE TABLE IF NOT EXISTS users (
            user_id INTEGER PRIMARY KEY, username TEXT, full_name TEXT,
            language TEXT DEFAULT 'uz', joined_at TEXT DEFAULT CURRENT_TIMESTAMP)""")
        await db.execute("""CREATE TABLE IF NOT EXISTS orders (
            id INTEGER PRIMARY KEY AUTOINCREMENT, user_id INTEGER,
            doc_type TEXT, amount INTEGER, status TEXT DEFAULT 'pending',
            data TEXT, file_path TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP)""")
        await db.commit()

async def save_user(user_id, username, full_name):
    async with aiosqlite.connect(DB) as db:
        await db.execute("INSERT OR IGNORE INTO users (user_id, username, full_name) VALUES (?,?,?)",
                         (user_id, username, full_name))
        await db.commit()

async def get_lang(user_id):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT language FROM users WHERE user_id=?", (user_id,)) as cur:
            row = await cur.fetchone()
            return row[0] if row else "uz"

async def set_lang(user_id, lang):
    async with aiosqlite.connect(DB) as db:
        await db.execute("UPDATE users SET language=? WHERE user_id=?", (lang, user_id))
        await db.commit()

async def create_order(user_id, doc_type, amount, data):
    async with aiosqlite.connect(DB) as db:
        cur = await db.execute(
            "INSERT INTO orders (user_id, doc_type, amount, data) VALUES (?,?,?,?)",
            (user_id, doc_type, amount, json.dumps(data, ensure_ascii=False)))
        await db.commit()
        return cur.lastrowid

async def get_order(order_id):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT * FROM orders WHERE id=?", (order_id,)) as cur:
            row = await cur.fetchone()
            if row:
                return {"id": row[0], "user_id": row[1], "doc_type": row[2],
                        "amount": row[3], "status": row[4],
                        "data": json.loads(row[5] or "{}"), "file_path": row[6]}

async def update_order(order_id, status, file_path=None):
    async with aiosqlite.connect(DB) as db:
        await db.execute("UPDATE orders SET status=?, file_path=? WHERE id=?",
                         (status, file_path, order_id))
        await db.commit()

async def get_stats():
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT COUNT(*) FROM users") as c:
            users = (await c.fetchone())[0]
        async with db.execute("SELECT COUNT(*), SUM(amount) FROM orders WHERE status='paid'") as c:
            row = await c.fetchone()
            orders, revenue = row[0], row[1] or 0
        async with db.execute("SELECT COUNT(*) FROM users WHERE joined_at >= date('now','-1 day')") as c:
            new_today = (await c.fetchone())[0]
    return users, orders, revenue, new_today

async def get_recent_users(limit=10):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT user_id, username, full_name, joined_at FROM users ORDER BY joined_at DESC LIMIT ?", (limit,)) as cur:
            return await cur.fetchall()

async def get_recent_orders(limit=10):
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT id, user_id, doc_type, amount, status, created_at FROM orders ORDER BY created_at DESC LIMIT ?", (limit,)) as cur:
            return await cur.fetchall()

async def get_all_users():
    async with aiosqlite.connect(DB) as db:
        async with db.execute("SELECT user_id FROM users") as cur:
            return [r[0] for r in await cur.fetchall()]

# ─── HUJJAT GENERATSIYA ───────────────────────────────────────────────────────

def make_doc(doc_type, data, order_id):
    # Avvalgi turlar
    if doc_type in ["ijara", "oldi_sotdi", "tilxat", "ishonchnoma", "nikoh", "qarz"]:
        return _original_make_doc(doc_type, data, order_id)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    sana = datetime.now().strftime('%d.%m.%Y')

    def heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text).bold = True

    def line(text):
        doc.add_paragraph(text)

    if doc_type == "mehnat":
        heading("MEHNAT SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"Ish beruvchi: {data.get('ish_beruvchi','___')}")
        line(f"Xodim: {data.get('xodim_fio','___')} (pasport: {data.get('xodim_pasport','___')})")
        line(f"Lavozim: {data.get('lavozim','___')}")
        line(f"Oylik maosh: {data.get('oylik','___')} so'm")
        line(f"Shartnoma muddati: {data.get('muddat','___')}")
        line(f"\nIsh vaqti: Dushanba-Juma, 09:00-18:00")
        line(f"Ta'til: Yiliga 15 ish kuni")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "pudrat":
        heading("PUDRAT SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"Buyurtmachi: {data.get('buyurtmachi','___')}")
        line(f"Pudratchi: {data.get('pudratchi','___')} ({data.get('pudratchi_pasport','___')})")
        line(f"Ish tavsifi: {data.get('ish_tavsifi','___')}")
        line(f"Ish narxi: {data.get('narx','___')} so'm")
        line(f"Bajarish muddati: {data.get('muddat','___')}")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "aliment":
        heading("ALIMENT TO'LASH SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"To'lovchi: {data.get('tolovchi','___')} (pasport: {data.get('tolovchi_pasport','___')})")
        line(f"Oluvchi: {data.get('oluvchi','___')}")
        line(f"Farzand(lar): {data.get('bola_ismi','___')}")
        line(f"Oylik miqdor: {data.get('miqdor','___')} so'm")
        line(f"Muddat: {data.get('muddat','___')}")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "hamkorlik":
        heading("HAMKORLIK SHARTNOMASI")
        line(f"Sana: {sana}  |  Shahar: {data.get('shahar','')}")
        line(f"1-tomon: {data.get('tomon1','___')} ({data.get('tomon1_pasport','___')})")
        line(f"2-tomon: {data.get('tomon2','___')} ({data.get('tomon2_pasport','___')})")
        line(f"Loyiha: {data.get('loyiha','___')}")
        line(f"Foyda taqsimoti: {data.get('foyda_taqsim','___')}")
        line(f"Muddat: {data.get('muddat','___')}")
        line("\nImzo: __________          Imzo: __________")

    elif doc_type == "davo":
        heading("DA'VO ARIZASI")
        line(f"Sana: {sana}")
        line(f"Sudga: {data.get('shahar','___')} tumani sudiga")
        line(f"\nAriza beruvchi: {data.get('ariza_beruvchi','___')} (pasport: {data.get('pasport','___')})")
        line(f"Javobgar: {data.get('javobgar','___')}")
        line(f"\nMuammo: {data.get('muammo','___')}")
        line(f"\nTalabim: {data.get('talab','___')}")
        line(f"\nAriza beruvchi imzosi: __________")
        line(f"Sana: {sana}")

    path = f"{DOCS_DIR}/{doc_type}_{order_id}.docx"
    doc.save(path)
    return path

# ─── YANGI HUJJATLAR MENYUSI ─────────────────────────────────────────────────

@dp.callback_query(F.data == "menu_docs")
async def docs_new(cb: CallbackQuery):
    lang = await get_user_lang(cb.from_user.id)
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📁 Asosiy hujjatlar", callback_data="cat_asosiy")],
        [InlineKeyboardButton(text="📝 Ish hujjatlari", callback_data="cat_ish")],
        [InlineKeyboardButton(text="🏗️ Qurilish", callback_data="cat_qurilish")],
        [InlineKeyboardButton(text="👨‍👩‍👧 Oila", callback_data="cat_oila")],
        [InlineKeyboardButton(text="🏢 Biznes", callback_data="cat_biznes")],
        [InlineKeyboardButton(text="⚖️ Yuridik", callback_data="cat_yuridik")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="back_main")],
    ])
    await cb.message.edit_text("📂 <b>Kategoriyani tanlang:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_asosiy")
async def cat_asosiy(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏠 Ijara shartnomasi", callback_data="doc_ijara")],
        [InlineKeyboardButton(text="🤝 Oldi-sotdi", callback_data="doc_oldi_sotdi")],
        [InlineKeyboardButton(text="✍️ Tilxat", callback_data="doc_tilxat")],
        [InlineKeyboardButton(text="📋 Ishonchnoma", callback_data="doc_ishonchnoma")],
        [InlineKeyboardButton(text="💍 Nikoh shartnomasi", callback_data="doc_nikoh")],
        [InlineKeyboardButton(text="💰 Qarz shartnomasi", callback_data="doc_qarz")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("📁 <b>Asosiy hujjatlar:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_ish")
async def cat_ish(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="📝 Mehnat shartnomasi", callback_data="doc_mehnat")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("📝 <b>Ish hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_qurilish")
async def cat_qurilish(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏗️ Pudrat shartnomasi", callback_data="doc_pudrat")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("🏗️ <b>Qurilish hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_oila")
async def cat_oila(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="👨‍👩‍👧 Aliment shartnomasi", callback_data="doc_aliment")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("👨‍👩‍👧 <b>Oila hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_biznes")
async def cat_biznes(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🏢 Hamkorlik shartnomasi", callback_data="doc_hamkorlik")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("🏢 <b>Biznes hujjatlari:</b>", parse_mode="HTML", reply_markup=keyboard)

@dp.callback_query(F.data == "cat_yuridik")
async def cat_yuridik(cb: CallbackQuery):
    keyboard = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="⚖️ Da'vo arizasi", callback_data="doc_davo")],
        [InlineKeyboardButton(text="⬅️ Orqaga", callback_data="menu_docs")],
    ])
    await cb.message.edit_text("⚖️ <b>Yuridik hujjatlar:</b>", parse_mode="HTML", reply_markup=keyboard)
