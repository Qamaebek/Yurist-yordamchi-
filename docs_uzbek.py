from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

DOCS_DIR = "docs"
os.makedirs(DOCS_DIR, exist_ok=True)

BLUE = RGBColor(27, 58, 107)
WHITE = RGBColor(255, 255, 255)
GRAY = RGBColor(150, 150, 150)
LIGHT_BLUE = 'DCE6F1'
DARK_BLUE = '1B3A6B'

def set_cell_color(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_border(table, color='1B3A6B'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for name in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tblBorders.append(b)
    tblPr.append(tblBorders)

def add_header(doc, title, qonun=''):
    # Ko'k sarlavha
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 1-qator: O'zbekiston Respublikasi
    c1 = t.cell(0, 0)
    set_cell_color(c1, DARK_BLUE)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("O'ZBEKISTON RESPUBLIKASI")
    r1.font.color.rgb = WHITE
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.name = 'Times New Roman'

    # 2-qator: Gerb o'rniga sarlavha
    c2 = t.cell(1, 0)
    set_cell_color(c2, 'EEF3FA')
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("YURIDIK YORDAMCHI XIZMATI  ⚖️")
    r2.font.color.rgb = BLUE
    r2.font.size = Pt(10)
    r2.font.italic = True
    r2.font.name = 'Times New Roman'

    doc.add_paragraph()

    # Asosiy sarlavha
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.font.size = Pt(16)
    tr.font.bold = True
    tr.font.name = 'Times New Roman'
    tr.font.color.rgb = BLUE

    if qonun:
        qp = doc.add_paragraph()
        qp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qr = qp.add_run(f"({qonun})")
        qr.font.size = Pt(9)
        qr.font.italic = True
        qr.font.name = 'Times New Roman'
        qr.font.color.rgb = GRAY

    # Ajratuvchi chiziq
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run("─" * 65)
    lr.font.color.rgb = BLUE
    lr.font.size = Pt(9)
    doc.add_paragraph()

def add_section_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

def add_table(doc, rows_data, header=None):
    total_rows = len(rows_data) + (1 if header else 0)
    table = doc.add_table(rows=total_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_border(table)

    offset = 0
    if header:
        hc = table.cell(0, 0)
        hc2 = table.cell(0, 1)
        hc.merge(hc2)
        set_cell_color(hc, DARK_BLUE)
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(header)
        hr.font.bold = True
        hr.font.color.rgb = WHITE
        hr.font.size = Pt(11)
        hr.font.name = 'Times New Roman'
        offset = 1

    for i, (label, value) in enumerate(rows_data):
        lc = table.cell(i + offset, 0)
        set_cell_color(lc, LIGHT_BLUE)
        lp = lc.paragraphs[0]
        lr = lp.add_run(f"  {label}")
        lr.font.bold = True
        lr.font.size = Pt(11)
        lr.font.name = 'Times New Roman'
        lr.font.color.rgb = BLUE

        vc = table.cell(i + offset, 1)
        vp = vc.paragraphs[0]
        vr = vp.add_run(f"  {value}")
        vr.font.size = Pt(11)
        vr.font.name = 'Times New Roman'

    doc.add_paragraph()

def add_clause(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

def add_signature(doc, left_title, right_title, left_name, right_name):
    doc.add_paragraph()
    # Sarlavha
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("TOMONLARNING REKVIZITLARI VA IMZOLARI")
    sr.font.bold = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = BLUE
    sr.font.name = 'Times New Roman'
    doc.add_paragraph()

    sana = datetime.now().strftime('%d.%m.%Y')

    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_border(table)

    # Sarlavhalar
    for col, title in enumerate([left_title, right_title]):
        cell = table.cell(0, col)
        set_cell_color(cell, DARK_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(11)
        r.font.name = 'Times New Roman'

    rows_data = [
        ("F.I.O.:", [left_name, right_name]),
        ("Pasport:", ["___________________", "___________________"]),
        ("Manzil:", ["___________________", "___________________"]),
        ("Tel:", ["___________________", "___________________"]),
        ("Imzo:", ["___________________", "___________________"]),
        ("Sana:", [sana, sana]),
    ]

    for i, (label, values) in enumerate(rows_data):
        for col in range(2):
            cell = table.cell(i + 1, col)
            if col == 0 or True:
                p = cell.paragraphs[0]
                if col == 0:
                    set_cell_color(cell, LIGHT_BLUE)
                    r = p.add_run(f"  {label}  {values[col]}")
                else:
                    r = p.add_run(f"  {label}  {values[col]}")
                r.font.size = Pt(11)
                r.font.name = 'Times New Roman'
                if col == 0:
                    r.font.color.rgb = BLUE

def add_footer(doc):
    doc.add_paragraph()
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run("─" * 65)
    lr.font.color.rgb = BLUE
    lr.font.size = Pt(9)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Ushbu hujjat yuridik kuchga ega bo'lib, O'zbekiston Respublikasi qonunchiligi asosida tuzilgan.")
    fr.font.size = Pt(8)
    fr.font.italic = True
    fr.font.color.rgb = GRAY
    fr.font.name = 'Times New Roman'
    fp2 = doc.add_paragraph()
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr2 = fp2.add_run("© Yuridik Yordamchi Bot | @YuridikYordamchi")
    fr2.font.size = Pt(8)
    fr2.font.italic = True
    fr2.font.color.rgb = GRAY
    fr2.font.name = 'Times New Roman'

def new_doc():
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    return doc

def make_doc(doc_type, data, order_id):
    doc = new_doc()
    sana = datetime.now().strftime('%d.%m.%Y')
    shahar = data.get('shahar', 'Toshkent')
    # "shahri shahri" muammosini tuzatish
    if shahar.endswith(' shahri'):
        shahar_full = shahar
    else:
        shahar_full = shahar + ' shahri'

    # ── IJARA ──────────────────────────────────────────────────────────────
    if doc_type == "ijara":
        add_header(doc, "KO'CHMAS MULKNI IJARAGA BERISH SHARTNOMASI",
                   "O'zbekiston Respublikasi Fuqarolik kodeksining 535-572-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"YY-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("Ijara beruvchi (F.I.O.):", data.get('beruvchi_fio', '___')),
            ("Pasport seriyasi:", data.get('beruvchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
            ("Ijara oluvchi (F.I.O.):", data.get('oluvchi_fio', '___')),
            ("Pasport seriyasi:", data.get('oluvchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
        ])

        add_section_title(doc, "2. SHARTNOMA PREDMETI")
        add_table(doc, [
            ("Mulk manzili:", data.get('manzil', '___')),
            ("Mulk turi:", "Ko'chmas mulk (turar joy/noturar joy)"),
            ("Ijara muddati:", f"{data.get('muddat', '___')} oy"),
            ("Boshlanish sanasi:", sana),
            ("Tugash sanasi:", "___________________"),
        ])

        add_section_title(doc, "3. IJARA HAQI VA TO'LOV TARTIBI")
        add_table(doc, [
            ("Oylik ijara haqi:", f"{data.get('narx', '___')} so'm"),
            ("So'zlar bilan:", "___________________"),
            ("To'lov muddati:", "Har oyning 5-kunigacha"),
            ("To'lov usuli:", "Naqd pul yoki bank o'tkazmasi"),
            ("Garov summasi:", "___________________"),
        ])

        add_section_title(doc, "4. TOMONLARNING HUQUQ VA MAJBURIYATLARI")
        add_clause(doc, "4.1. Ijara beruvchi mulkni yaxshi texnik holatda topshirishga majbur.")
        add_clause(doc, "4.2. Ijara oluvchi mulkni faqat maqsadli foydalanishga majbur.")
        add_clause(doc, "4.3. Ijara oluvchi mulkni uchinchi shaxslarga ijaraga berish huquqiga ega emas.")
        add_clause(doc, "4.4. Kommunal to'lovlar (gaz, suv, elektr) ijara oluvchi tomonidan to'lanadi.")
        add_clause(doc, "4.5. Joriy ta'mirlash ijara oluvchi, kapital ta'mirlash ijara beruvchi tomonidan amalga oshiriladi.")
        doc.add_paragraph()

        add_section_title(doc, "5. SHARTNOMANI BEKOR QILISH TARTIBI")
        add_clause(doc, "5.1. Har bir tomon shartnomani bir oy oldin xabardor qilgan holda bekor qilish huquqiga ega.")
        add_clause(doc, "5.2. Ijara haqi to'lanmagan taqdirda, ijara beruvchi shartnomani bir tomonlama bekor qilish huquqiga ega.")
        doc.add_paragraph()

        add_section_title(doc, "6. NIZOLARNI HAL ETISH TARTIBI")
        add_clause(doc, f"6.1. Nizolar muzokaralar yo'li bilan hal etiladi.")
        add_clause(doc, f"6.2. Kelishuv bo'lmagan taqdirda, nizolar {shahar_full} sudi orqali hal etiladi.")
        doc.add_paragraph()

        add_section_title(doc, "7. SHARTNOMANING AMAL QILISH MUDDATI")
        add_clause(doc, f"7.1. Shartnoma imzolangan kundan boshlab {data.get('muddat', '___')} oy davomida amal qiladi.")
        add_clause(doc, "7.2. Shartnoma 2 nusxada tuzilgan bo'lib, har bir tomon uchun bir nusxa beriladi.")
        doc.add_paragraph()

        add_signature(doc, "IJARA BERUVCHI", "IJARA OLUVCHI",
                     data.get('beruvchi_fio', '___'), data.get('oluvchi_fio', '___'))

    # ── OLDI-SOTDI ─────────────────────────────────────────────────────────
    elif doc_type == "oldi_sotdi":
        add_header(doc, "OLDI-SOTDI SHARTNOMASI",
                   "O'zbekiston Respublikasi Fuqarolik kodeksining 386-425-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"OS-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("Sotuvchi (F.I.O.):", data.get('sotuvchi_fio', '___')),
            ("Pasport seriyasi:", data.get('sotuvchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
            ("Xaridor (F.I.O.):", data.get('xaridor_fio', '___')),
            ("Pasport seriyasi:", data.get('xaridor_pasport', '___')),
            ("Yashash manzili:", "___________________"),
        ])

        add_section_title(doc, "2. SHARTNOMA PREDMETI")
        add_table(doc, [
            ("Mulk tavsifi:", data.get('tovar', '___')),
            ("Mulk holati:", "Ishlatilgan / Yangi"),
            ("Mulk narxi:", f"{data.get('narx', '___')} so'm"),
            ("So'zlar bilan:", "___________________"),
        ])

        add_section_title(doc, "3. TO'LOV TARTIBI")
        add_clause(doc, "3.1. Xaridor mulk narxini shartnoma imzolanish kuni to'liq to'laydi.")
        add_clause(doc, "3.2. Mulk huquqi xaridorga to'lov amalga oshirilganidan so'ng o'tadi.")
        add_clause(doc, "3.3. To'lov naqd pul yoki bank o'tkazmasi orqali amalga oshiriladi.")
        doc.add_paragraph()

        add_section_title(doc, "4. TOMONLARNING MAJBURIYATLARI")
        add_clause(doc, "4.1. Sotuvchi mulkni belgilangan holatda va uchinchi shaxslarning da'vosidan xoli holda topshiradi.")
        add_clause(doc, "4.2. Xaridor mulk narxini to'liq to'lashga majbur.")
        add_clause(doc, "4.3. Sotuvchi mulkda yashirin nuqsonlar bo'lsa, ularni xaridorga oldindan xabar qiladi.")
        doc.add_paragraph()

        add_section_title(doc, "5. KAFOLAT VA JAVOBGARLIK")
        add_clause(doc, "5.1. Sotuvchi mulkda yashirin nuqsonlar uchun 6 oy davomida javobgar hisoblanadi.")
        add_clause(doc, "5.2. Shartnoma bekor qilinsa, tomonlar dastlabki holatga qaytariladi.")
        doc.add_paragraph()

        add_section_title(doc, "6. NIZOLARNI HAL ETISH")
        add_clause(doc, f"6.1. Nizolar {shahar_full} sudi orqali hal etiladi.")

        add_signature(doc, "SOTUVCHI", "XARIDOR",
                     data.get('sotuvchi_fio', '___'), data.get('xaridor_fio', '___'))

    # ── TILXAT ─────────────────────────────────────────────────────────────
    elif doc_type == "tilxat":
        add_header(doc, "TILXAT",
                   "O'zbekiston Respublikasi Fuqarolik kodeksining 732-747-moddalari asosida")

        add_table(doc, [
            ("Yozilgan sana:", sana),
            ("Yozilgan joy:", shahar_full),
        ])

        add_section_title(doc, "TILXAT YOZUVCHINING MA'LUMOTLARI")
        add_table(doc, [
            ("F.I.O.:", data.get('fio', '___')),
            ("Pasport seriyasi:", data.get('pasport', '___')),
            ("Yashash manzili:", data.get('manzil', '___')),
            ("Telefon:", "___________________"),
        ])

        add_section_title(doc, "TILXAT MAZMUNI")
        doc.add_paragraph()
        tp = doc.add_paragraph()
        tp.paragraph_format.left_indent = Cm(1)
        tp.paragraph_format.right_indent = Cm(1)
        tr = tp.add_run(f"Men, {data.get('fio', '___')}, ushbu tilxat bilan tasdiqlayman:")
        tr.font.size = Pt(12)
        tr.font.name = 'Times New Roman'
        doc.add_paragraph()

        mp = doc.add_paragraph()
        mp.paragraph_format.left_indent = Cm(1)
        mp.paragraph_format.right_indent = Cm(1)
        mr = mp.add_run(data.get('mazmun', '___'))
        mr.font.size = Pt(12)
        mr.font.bold = True
        mr.font.name = 'Times New Roman'
        doc.add_paragraph()

        add_table(doc, [
            ("Miqdori:", f"{data.get('miqdor', '___')} so'm"),
            ("So'zlar bilan:", "___________________"),
            ("Qaytarish muddati:", "___________________"),
        ])

        add_clause(doc, "Ushbu tilxat O'zbekiston Respublikasi Fuqarolik kodeksi 732-moddasiga muvofiq tuzilgan va yuridik kuchga ega.")
        doc.add_paragraph()

        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ir = ip.add_run(f"F.I.O.: {data.get('fio', '___')}\nImzo: ___________________\nSana: {sana}")
        ir.font.size = Pt(12)
        ir.font.name = 'Times New Roman'

    # ── ISHONCHNOMA ────────────────────────────────────────────────────────
    elif doc_type == "ishonchnoma":
        add_header(doc, "ISHONCHNOMA",
                   "O'zbekiston Respublikasi Fuqarolik kodeksining 121-129-moddalari asosida")

        add_table(doc, [
            ("Raqam:", f"ISH-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "ISHONCHNOMA BERUVCHI")
        add_table(doc, [
            ("F.I.O.:", data.get('beruvchi_fio', '___')),
            ("Pasport seriyasi:", data.get('beruvchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
            ("Telefon:", "___________________"),
        ])

        add_section_title(doc, "ISHONCHNOMA OLUVCHI (VAKIL)")
        add_table(doc, [
            ("F.I.O.:", data.get('oluvchi_fio', '___')),
            ("Pasport seriyasi:", data.get('oluvchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
        ])

        add_section_title(doc, "BERILGAN VAKOLATLAR")
        vp = doc.add_paragraph()
        vp.paragraph_format.left_indent = Cm(1)
        vr = vp.add_run(data.get('vakolat', '___'))
        vr.font.size = Pt(12)
        vr.font.name = 'Times New Roman'
        doc.add_paragraph()

        add_table(doc, [
            ("Amal qilish muddati:", data.get('muddat', '___')),
            ("Boshlanish sanasi:", sana),
            ("Qayta topshirish:", "Taqiqlanadi"),
        ])

        add_clause(doc, "Ushbu ishonchnoma notarius tomonidan tasdiqlanishi tavsiya etiladi.")

        add_signature(doc, "ISHONCHNOMA BERUVCHI", "VAKIL",
                     data.get('beruvchi_fio', '___'), data.get('oluvchi_fio', '___'))

    # ── NIKOH ──────────────────────────────────────────────────────────────
    elif doc_type == "nikoh":
        add_header(doc, "NIKOH SHARTNOMASI",
                   "O'zbekiston Respublikasi Oila kodeksining 27-31-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"NK-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("Er (F.I.O.):", data.get('er_fio', '___')),
            ("Pasport seriyasi:", data.get('er_pasport', '___')),
            ("Tug'ilgan sanasi:", "___________________"),
            ("Xotin (F.I.O.):", data.get('xotin_fio', '___')),
            ("Pasport seriyasi:", data.get('xotin_pasport', '___')),
            ("Tug'ilgan sanasi:", "___________________"),
        ])

        add_section_title(doc, "2. NIKOHGACHA BO'LGAN MULK")
        add_clause(doc, "2.1. Nikohgacha har bir tomonning shaxsiy mulki hisoblanadi:")
        add_clause(doc, "Er: ___________________", indent=True)
        add_clause(doc, "Xotin: ___________________", indent=True)
        doc.add_paragraph()

        add_section_title(doc, "3. NIKOH DAVOMIDA MULK")
        add_clause(doc, f"3.1. Mulkiy shartlar: {data.get('shartlar', '___')}")
        add_clause(doc, "3.2. Nikoh davomida orttirilgan mulk qonun bo'yicha umumiy hisoblanadi.")
        add_clause(doc, "3.3. Meros va hadya orqali olingan mulk shaxsiy mulk hisoblanadi.")
        doc.add_paragraph()

        add_section_title(doc, "4. UMUMIY QOIDALAR")
        add_clause(doc, "4.1. Shartnoma notarius tomonidan tasdiqlanganidan keyin kuchga kiradi (OK 29-modda).")
        add_clause(doc, "4.2. Shartnomaga o'zgartirishlar faqat tomonlarning kelishuvi bilan kiritiladi.")
        add_clause(doc, "4.3. Ushbu shartnoma O'zbekiston Respublikasi Oila kodeksiga zid kelmasligi shart.")

        add_signature(doc, "ER", "XOTIN",
                     data.get('er_fio', '___'), data.get('xotin_fio', '___'))

    # ── QARZ ───────────────────────────────────────────────────────────────
    elif doc_type == "qarz":
        add_header(doc, "QARZ SHARTNOMASI",
                   "O'zbekiston Respublikasi Fuqarolik kodeksining 732-747-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"QR-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("Qarz beruvchi (F.I.O.):", data.get('beruvchi_fio', '___')),
            ("Pasport seriyasi:", data.get('beruvchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
            ("Qarz oluvchi (F.I.O.):", data.get('oluvchi_fio', '___')),
            ("Pasport seriyasi:", data.get('oluvchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
        ])

        add_section_title(doc, "2. QARZ SHARTLARI")
        add_table(doc, [
            ("Qarz miqdori:", f"{data.get('miqdor', '___')} so'm"),
            ("So'zlar bilan:", "___________________"),
            ("Foiz stavkasi:", f"{data.get('foiz', '0')}% yillik"),
            ("Qaytarish muddati:", data.get('muddat', '___')),
            ("Qaytarish sanasi:", "___________________"),
            ("Qaytarish usuli:", "Naqd pul yoki bank o'tkazmasi"),
        ])

        add_section_title(doc, "3. TOMONLARNING MAJBURIYATLARI")
        add_clause(doc, "3.1. Qarz beruvchi kelishilgan miqdorni shartnoma imzolanishi bilan beradi.")
        add_clause(doc, "3.2. Qarz oluvchi qarzni belgilangan muddat va usulda to'liq qaytaradi.")
        add_clause(doc, "3.3. Muddatidan kech qaytarilganda qarz oluvchi har kuni 0.1% jarima to'laydi.")
        add_clause(doc, "3.4. Qarz beruvchi mablag'ni maqsadli sarflanishini tekshirish huquqiga ega.")
        doc.add_paragraph()

        add_section_title(doc, "4. NIZOLARNI HAL ETISH")
        add_clause(doc, "4.1. Nizolar avval muzokaralar yo'li bilan hal etiladi.")
        add_clause(doc, f"4.2. Kelishuv bo'lmagan taqdirda {shahar_full} sudi orqali hal etiladi.")
        doc.add_paragraph()

        add_section_title(doc, "5. YAKUNIY QOIDALAR")
        add_clause(doc, "5.1. Shartnoma 2 nusxada tuzilgan, har bir tomon uchun 1 nusxa.")
        add_clause(doc, "5.2. Shartnomaga o'zgartirishlar yozma kelishuv bilan kiritiladi.")

        add_signature(doc, "QARZ BERUVCHI", "QARZ OLUVCHI",
                     data.get('beruvchi_fio', '___'), data.get('oluvchi_fio', '___'))

    # ── MEHNAT ─────────────────────────────────────────────────────────────
    elif doc_type == "mehnat":
        add_header(doc, "MEHNAT SHARTNOMASI",
                   "O'zbekiston Respublikasi Mehnat kodeksining 73-100-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"MH-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("Ish beruvchi:", data.get('ish_beruvchi', '___')),
            ("STIR/INN:", "___________________"),
            ("Manzil:", "___________________"),
            ("Xodim (F.I.O.):", data.get('xodim_fio', '___')),
            ("Pasport seriyasi:", data.get('xodim_pasport', '___')),
            ("Yashash manzili:", "___________________"),
        ])

        add_section_title(doc, "2. ISH SHAROITLARI")
        add_table(doc, [
            ("Lavozim:", data.get('lavozim', '___')),
            ("Bo'lim:", "___________________"),
            ("Oylik maosh:", f"{data.get('oylik', '___')} so'm"),
            ("Mukofot:", "Ish natijasiga qarab"),
            ("Shartnoma turi:", "Muddatli / Muddatsiz"),
            ("Shartnoma muddati:", data.get('muddat', '___')),
            ("Sinov muddati:", "3 oy (MK 84-modda)"),
        ])

        add_section_title(doc, "3. ISH VAQTI VA DAM OLISH")
        add_table(doc, [
            ("Ish vaqti:", "09:00 - 18:00"),
            ("Ish kunlari:", "Dushanba - Juma"),
            ("Tanaffus:", "13:00 - 14:00 (1 soat)"),
            ("Dam olish kunlari:", "Shanba, Yakshanba"),
            ("Yillik ta'til:", "15 ish kuni (MK 134-modda)"),
        ])

        add_section_title(doc, "4. TOMONLARNING MAJBURIYATLARI")
        add_clause(doc, "4.1. Ish beruvchi maoshni har oyning oxirgi ish kunigacha to'lashga majbur.")
        add_clause(doc, "4.2. Ish beruvchi xodimni ijtimoiy sug'urta qilishga majbur.")
        add_clause(doc, "4.3. Xodim ish joyida intizom va maxfiylikka rioya qilishga majbur.")
        add_clause(doc, "4.4. Shartnomani bekor qilishda 2 hafta oldindan xabardor qilish shart (MK 97-modda).")
        doc.add_paragraph()

        add_section_title(doc, "5. NIZOLARNI HAL ETISH")
        add_clause(doc, f"5.1. Mehnat nizolari {shahar_full} mehnat inspeksiyasi yoki sudi orqali hal etiladi.")

        add_signature(doc, "ISH BERUVCHI", "XODIM",
                     data.get('ish_beruvchi', '___'), data.get('xodim_fio', '___'))

    # ── PUDRAT ─────────────────────────────────────────────────────────────
    elif doc_type == "pudrat":
        add_header(doc, "PUDRAT SHARTNOMASI",
                   "O'zbekiston Respublikasi Fuqarolik kodeksining 631-665-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"PD-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("Buyurtmachi (F.I.O./Tashkilot):", data.get('buyurtmachi', '___')),
            ("Pasport/STIR:", "___________________"),
            ("Manzil:", "___________________"),
            ("Pudratchi (F.I.O./Tashkilot):", data.get('pudratchi', '___')),
            ("Pasport/STIR:", data.get('pudratchi_pasport', '___')),
            ("Manzil:", "___________________"),
        ])

        add_section_title(doc, "2. ISH PREDMETI VA NARXI")
        add_table(doc, [
            ("Bajariladigan ish:", data.get('ish_tavsifi', '___')),
            ("Ish joyi:", "___________________"),
            ("Ish narxi:", f"{data.get('narx', '___')} so'm"),
            ("So'zlar bilan:", "___________________"),
            ("Bajarish muddati:", data.get('muddat', '___')),
            ("Topshirish sanasi:", "___________________"),
        ])

        add_section_title(doc, "3. TO'LOV TARTIBI")
        add_clause(doc, "3.1. Avans: Shartnoma imzolanishi bilan ish narxining 30% to'lanadi.")
        add_clause(doc, "3.2. Qolgan 70%: Ish qabul qilinib, dalolatnoma imzolangandan so'ng to'lanadi.")
        add_clause(doc, "3.3. Kechikish uchun kunlik 0.1% jarima qo'llaniladi.")
        doc.add_paragraph()

        add_section_title(doc, "4. KAFOLAT VA SIFAT")
        add_clause(doc, "4.1. Pudratchi bajarilgan ish sifatiga 12 oy kafolat beradi.")
        add_clause(doc, "4.2. Kafolat muddatida kamchiliklar pudratchi hisobiga bartaraf etiladi.")
        doc.add_paragraph()

        add_section_title(doc, "5. NIZOLARNI HAL ETISH")
        add_clause(doc, f"5.1. Nizolar {shahar_full} sudi orqali hal etiladi.")

        add_signature(doc, "BUYURTMACHI", "PUDRATCHI",
                     data.get('buyurtmachi', '___'), data.get('pudratchi', '___'))

    # ── ALIMENT ────────────────────────────────────────────────────────────
    elif doc_type == "aliment":
        add_header(doc, "ALIMENT TO'LASH TO'G'RISIDA SHARTNOMA",
                   "O'zbekiston Respublikasi Oila kodeksining 99-113-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"AL-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("Aliment to'lovchi (F.I.O.):", data.get('tolovchi', '___')),
            ("Pasport seriyasi:", data.get('tolovchi_pasport', '___')),
            ("Yashash manzili:", "___________________"),
            ("Aliment oluvchi (F.I.O.):", data.get('oluvchi', '___')),
            ("Yashash manzili:", "___________________"),
        ])

        add_section_title(doc, "2. FARZAND MA'LUMOTLARI")
        add_table(doc, [
            ("Farzand(lar) F.I.O.:", data.get('bola_ismi', '___')),
            ("Tug'ilgan sanasi:", "___________________"),
            ("Yashash manzili:", "___________________"),
        ])

        add_section_title(doc, "3. ALIMENT MIQDORI VA TO'LOV TARTIBI")
        add_table(doc, [
            ("Oylik aliment miqdori:", f"{data.get('miqdor', '___')} so'm"),
            ("So'zlar bilan:", "___________________"),
            ("To'lov sanasi:", "Har oyning 1-kunigacha"),
            ("To'lov usuli:", "Naqd pul yoki bank o'tkazmasi"),
            ("Amal qilish muddati:", data.get('muddat', '___')),
        ])

        add_section_title(doc, "4. UMUMIY QOIDALAR")
        add_clause(doc, "4.1. Shartnoma notarius tomonidan tasdiqlanishi shart (OK 101-modda).")
        add_clause(doc, "4.2. Kechikish holatida to'lovchi har kuni 0.5% ustama to'laydi.")
        add_clause(doc, "4.3. Aliment miqdori yiliga bir marta inflyatsiya darajasiga qarab indeksatsiya qilinadi.")
        add_clause(doc, "4.4. Farzand voyaga yetgach (18 yosh) to'lovlar avtomatik to'xtatiladi.")
        doc.add_paragraph()

        add_section_title(doc, "5. NIZOLARNI HAL ETISH")
        add_clause(doc, f"5.1. Nizolar {shahar_full} oila sudi orqali hal etiladi.")

        add_signature(doc, "ALIMENT TO'LOVCHI", "ALIMENT OLUVCHI",
                     data.get('tolovchi', '___'), data.get('oluvchi', '___'))

    # ── HAMKORLIK ──────────────────────────────────────────────────────────
    elif doc_type == "hamkorlik":
        add_header(doc, "HAMKORLIK (SHERIKLIK) SHARTNOMASI",
                   "O'zbekiston Respublikasi Fuqarolik kodeksining 781-796-moddalari asosida")

        add_table(doc, [
            ("Shartnoma №:", f"HK-{order_id}/{datetime.now().year}"),
            ("Sana:", sana),
            ("Joy:", shahar_full),
        ])

        add_section_title(doc, "1. SHARTNOMA TOMONLARI")
        add_table(doc, [
            ("1-tomon (F.I.O./Tashkilot):", data.get('tomon1', '___')),
            ("Pasport/STIR:", data.get('tomon1_pasport', '___')),
            ("Manzil:", "___________________"),
            ("2-tomon (F.I.O./Tashkilot):", data.get('tomon2', '___')),
            ("Pasport/STIR:", data.get('tomon2_pasport', '___')),
            ("Manzil:", "___________________"),
        ])

        add_section_title(doc, "2. HAMKORLIK MAQSADI VA PREDMETI")
        add_clause(doc, data.get('loyiha', '___'))
        doc.add_paragraph()

        add_section_title(doc, "3. TOMONLARNING ULUSHI VA FOYDA TAQSIMOTI")
        add_table(doc, [
            ("Foyda taqsimoti:", data.get('foyda_taqsim', '___')),
            ("Zarar taqsimoti:", data.get('foyda_taqsim', '___')),
            ("Hamkorlik muddati:", data.get('muddat', '___')),
            ("Moliyaviy hisob:", "Har oyning oxirida"),
        ])

        add_section_title(doc, "4. TOMONLARNING MAJBURIYATLARI")
        add_clause(doc, "4.1. Har bir tomon o'z ulushiga mos hissa qo'shadi.")
        add_clause(doc, "4.2. Muhim qarorlar tomonlarning birgalikda kelishuvi bilan qabul qilinadi.")
        add_clause(doc, "4.3. Moliyaviy hisobot har oy taqdim etiladi.")
        doc.add_paragraph()

        add_section_title(doc, "5. MAXFIYLIK")
        add_clause(doc, "5.1. Tomonlar hamkorlik doirasidagi ma'lumotlarni uchinchi shaxslarga oshkor etmaydi.")
        add_clause(doc, "5.2. Maxfiylik majburiyati shartnoma tugagandan keyin 2 yil davomida kuchda qoladi.")
        doc.add_paragraph()

        add_section_title(doc, "6. NIZOLARNI HAL ETISH")
        add_clause(doc, f"6.1. Nizolar {shahar_full} sudi orqali hal etiladi.")

        add_signature(doc, "1-TOMON", "2-TOMON",
                     data.get('tomon1', '___'), data.get('tomon2', '___'))

    # ── DA'VO ARIZASI ──────────────────────────────────────────────────────
    elif doc_type == "davo":
        add_header(doc, "DA'VO ARIZASI",
                   "O'zbekiston Respublikasi FPK 3-bo'limi asosida")

        jp = doc.add_paragraph()
        jp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        jr = jp.add_run(f"{shahar_full} fuqarolik sudi raisiga\nAriza beruvchi: {data.get('ariza_beruvchi', '___')}\nSana: {sana}")
        jr.font.size = Pt(11)
        jr.font.name = 'Times New Roman'
        doc.add_paragraph()

        add_section_title(doc, "1. ARIZA BERUVCHI MA'LUMOTLARI")
        add_table(doc, [
            ("F.I.O.:", data.get('ariza_beruvchi', '___')),
            ("Pasport seriyasi:", data.get('pasport', '___')),
            ("Yashash manzili:", "___________________"),
            ("Telefon:", "___________________"),
        ])

        add_section_title(doc, "2. JAVOBGAR MA'LUMOTLARI")
        add_table(doc, [
            ("Javobgar:", data.get('javobgar', '___')),
            ("Manzil:", "___________________"),
        ])

        add_section_title(doc, "3. ARIZANING MOHIYATI")
        mp = doc.add_paragraph()
        mp.paragraph_format.left_indent = Cm(1)
        mr = mp.add_run(data.get('muammo', '___'))
        mr.font.size = Pt(12)
        mr.font.name = 'Times New Roman'
        doc.add_paragraph()

        add_section_title(doc, "4. HUQUQIY ASOS")
        add_clause(doc, "O'zbekiston Respublikasi Fuqarolik kodeksi va Fuqarolik protsessual kodeksi asosida.")
        doc.add_paragraph()

        add_section_title(doc, "5. TALABLAR")
        tp = doc.add_paragraph()
        tp.paragraph_format.left_indent = Cm(1)
        tr = tp.add_run(data.get('talab', '___'))
        tr.font.size = Pt(12)
        tr.font.name = 'Times New Roman'
        doc.add_paragraph()

        add_section_title(doc, "6. ILOVALAR")
        add_clause(doc, "1. Ariza nusxasi — 2 dona")
        add_clause(doc, "2. Pasport nusxasi — 1 dona")
        add_clause(doc, "3. Dalil hujjatlar — _____ dona")
        add_clause(doc, "4. Davlat boji to'lovi cheki — 1 dona")
        doc.add_paragraph()

        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ir = ip.add_run(f"Ariza beruvchi: ___________________\nF.I.O.: {data.get('ariza_beruvchi', '___')}\nSana: {sana}")
        ir.font.size = Pt(12)
        ir.font.name = 'Times New Roman'

    add_footer(doc)
    path = f"{DOCS_DIR}/{doc_type}_{order_id}.docx"
    doc.save(path)
    return path
